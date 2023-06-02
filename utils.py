import re
import PyPDF2
import docx

from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# MAX_FILE_SIZE_MB = 2  # Maximum file size in megabytes

def preprocess_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        pages = []
        for page in pdf_reader.pages:
            pages.append(page.extract_text())
        return pages

def preprocess_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        texts = file.readlines()
        return texts

def preprocess_docx(file_path):
    doc = docx.Document(file_path)
    parag_texts = []
    for para in doc.paragraphs:
        parag_texts.append(para.text)
    return parag_texts

def merge_hyphenated_words(text: str) -> str:
    return re.sub(r"(\w)-\n(\w)", r"\1\2", text)


def fix_newlines(text: str) -> str:
    return re.sub(r"(?<!\n)\n(?!\n)", " ", text)


def remove_multiple_newlines(text: str) -> str:
    return re.sub(r"\n{2,}", "\n", text)


def clean_text(texts, cleaning_functions):
    cleaned_text = []
    for text in texts:
        for cleaning_function in cleaning_functions:
            text = cleaning_function(text)
        cleaned_text.append(text)
    return cleaned_text

def preprocess_file(file_path):
    try:
        if file_path.endswith('.pdf'):
            texts = preprocess_pdf(file_path)
        elif file_path.endswith('.txt'):
            texts = preprocess_txt(file_path)
        elif file_path.endswith('.docx'):
            texts = preprocess_docx(file_path)
        
        cleaning_functions = [
            merge_hyphenated_words,
            fix_newlines,
            remove_multiple_newlines,
        ]
        cleaned_text = clean_text(texts, cleaning_functions)
        doc_chunks = []

        for text in cleaned_text:
            text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=200,
            )

            chunks = text_splitter.split_text(text)
    
            for chunk in chunks:
                doc = Document(
                    page_content=chunk
                )

                doc_chunks.append(doc)

        return doc_chunks
    
    except ValueError as e:
        raise e("Unsupported file format. Accepted formats are .pdf, .docx and .txt")
    